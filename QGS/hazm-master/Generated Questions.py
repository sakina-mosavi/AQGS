from docx import Document
document = Document()

doc1 = document.add_heading("سوال صحیح و غلط")
doc1.alignment = 1

doc = document.add_paragraph("donedone")
doc.alignment = 2
doc = document.add_paragraph("nononoe")
doc.alignment = 2
doc = document.add_paragraph("tontone")
doc.alignment = 2
document.save("TF Generated Questions.docx")


