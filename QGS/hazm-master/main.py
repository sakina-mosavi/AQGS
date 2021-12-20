from __future__ import unicode_literals
####################################### accessing data or text ###############################################
import ImportData
# text = ImportData.text     # the text is in string format
text="افغانستان در قرن ۲۱ عضویت سازمان ملل متحد را بدست اورد."
import APISetup
import json
####################################### preprocessing step ################################################
import preprocessing
preprocessed = preprocessing.preprocessing(text)   # preprocessed type is list of sentence
# print(preprocessed)
# print(list(preprocessed))
# for W in preprocessed:
#     print(type(W))
############################################# Creating a .docx file  for generated questions #################################
from docx import Document
document = Document()

doc1 = document.add_heading("سوال صحیح و غلط")
doc1.alignment = 1

######################################################## TF questions ############################################################
import TFQuestion

import TFQuestion
for item in preprocessed:
    if (TFQuestion.changing_int_value_TF(item)):
        TFGeneratedQ = TFQuestion.changing_int_value_TF(item)
        doc = document.add_paragraph(TFGeneratedQ)
        doc.alignment = 2
    if (TFQuestion.Any_TF(item)):
        PersonQ = TFQuestion.Any_TF(item)
        doc = document.add_paragraph(PersonQ)
        doc.alignment = 2


def remove_brackets(data):
    ret = ''
    skip1c = 0
    skip2c = 0
    for i in data:
        if i == '[':
            skip1c += 1
        elif i == '(':
            skip2c += 1
        elif i == ']' and skip1c > 0:
            skip1c -= 1
        elif i == ')'and skip2c > 0:
            skip2c -= 1
        elif skip1c == 0 and skip2c == 0:
            ret += i
    return ret


doc = document.add_heading("سوالات خانه خالی")
doc.alignment = 1
###################################################### Blank Space Questions #########################################################
import BSQuestion
for item in preprocessed:
    if (BSQuestion.Int_BS(item)):
        doc = document.add_paragraph(BSQuestion.Int_BS(item))
        doc.alignment = 2
    if (BSQuestion.Per_BS(item)):
        item = remove_brackets(item)
        doc = document.add_paragraph(BSQuestion.Per_BS(item))
        doc.alignment = 2
    if ( BSQuestion.Dat_BS(item)):
        doc = document.add_paragraph(BSQuestion.Dat_BS(item))
        doc.alignment = 2
    if (BSQuestion.Loc_BS(item)):
        doc = document.add_paragraph(BSQuestion.Loc_BS(item))
        doc.alignment = 2
    if (BSQuestion.Eve_BS(item)):
        doc = document.add_paragraph(BSQuestion.Eve_BS(item))
        doc.alignment = 2
document.save("sixth.docx")
##################################################### Descriptive Question ################################################################
import DQuestion
doc = document.add_heading("سوالات تشریحی کوتاه")
doc.alignment = 1

for item in preprocessed:
    if(DQuestion.Int_D(item)):
        doc = document.add_paragraph(DQuestion.Int_D(item))
        doc.alignment = 2
    if (DQuestion.Person_Noun(item)):
        item = remove_brackets(item)
        doc = document.add_paragraph(DQuestion.Person_Noun(item))
        doc.alignment = 2
    if (DQuestion.Per_D(item)):
        doc = document.add_paragraph(DQuestion.Per_D(item))
        doc.alignment = 2
    if (DQuestion.Location_Noun(item)):
        doc = document.add_paragraph(DQuestion.Location_Noun(item))
        doc.alignment = 2
    if (DQuestion.Date_Noun(item)):
        doc = document.add_paragraph(DQuestion.Date_Noun(item))
        doc.alignment = 2
    if (DQuestion.Date_D(item)):
        doc = document.add_paragraph(DQuestion.Date_D(item))
        doc.alignment = 2
    if (DQuestion.Event(item)):
        doc = document.add_paragraph(DQuestion.Event(item))
        doc.alignment = 2

document.save("seventh.docx")
# "روزنامه سیمرغ در سال ۱۳۵۴ تاسیس شد.
# baseUrl = "http://api.text-mining.ir/api/"
# url =  baseUrl + "NamedEntityRecognition/Detect"
# payload = text
# tokenKey = APISetup.tokenkey()
# result = json.loads(APISetup.callApi(url, payload, tokenKey))
# for phrase in result:
#     print("("+phrase['word']+","+phrase['tags']['NER']['item1']+") ")
